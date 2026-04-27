"""
DOCX Generator — Creates lab experiment Word documents using python-docx.
"""
import io
import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def _set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    elem = shading.makeelement(qn("w:shd"), {qn("w:fill"): color_hex, qn("w:val"): "clear"})
    shading.append(elem)


def _strip_code_fences(code):
    code = code.strip()
    if code.startswith("```"):
        lines = code.split("\n")
        lines = lines[1:] if lines[0].startswith("```") else lines
        if lines and lines[-1].strip() == "```":
            lines = lines[:-1]
        code = "\n".join(lines)
    return code


def generate_experiment_docx(experiment: dict) -> io.BytesIO:
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2)

    exp_no = experiment.get("experiment_no", "")
    title = doc.add_heading(f"EXPERIMENT NO. {exp_no}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        run.font.size = Pt(18)

    date_str = datetime.datetime.now().strftime("%d/%m/%Y")
    date_p = doc.add_paragraph()
    date_run = date_p.add_run(f"DATE: {date_str}")
    date_run.font.bold = True
    date_run.font.size = Pt(11)

    def add_section(name, content, is_code=False):
        h = doc.add_heading(name, level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x16, 0x21, 0x3e)
            run.font.size = Pt(14)
        p = doc.add_paragraph()
        if is_code:
            content = _strip_code_fences(content)
        r = p.add_run(content)
        r.font.size = Pt(9 if is_code else 11)
        if is_code:
            r.font.name = "Courier New"
        return p

    add_section("AIM", experiment.get("aim", ""))
    add_section("THEORY", experiment.get("theory", ""))
    add_section("SOURCE CODE", experiment.get("source_code", ""), is_code=True)

    # VIVA VOCE
    h = doc.add_heading("VIVA VOCE", level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x16, 0x21, 0x3e)
        run.font.size = Pt(14)

    viva = experiment.get("viva", [])
    if viva:
        for idx, qa in enumerate(viva, 1):
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f"Q{idx}: {qa.get('question', '')}")
            q_run.font.bold = True
            
            a_para = doc.add_paragraph()
            a_run = a_para.add_run(f"A: {qa.get('answer', '')}")

    add_section("OUTPUT", experiment.get("output", ""), is_code=True)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
